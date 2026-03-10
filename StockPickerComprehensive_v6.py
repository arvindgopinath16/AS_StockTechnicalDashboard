import yfinance as yf
import pandas as pd
from ta.momentum import RSIIndicator
from ta.trend import MACD
from ta.volatility import BollingerBands
from ta.volume import VolumePriceTrendIndicator
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import time
import numpy as np
from yfinance import ticker


# Function to calculate VFI
def calculate_vfi(close, high, low, volume, window=100):
    mfv = ((close - low) - (high - close)) / (high - low) * volume
    vfi = mfv.rolling(window=window).sum()
    return vfi


# Function to create a new page in the document
def add_new_page(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


# Function to create summary table for BB, RSI, and MACD
def create_summary_table(document, stocks_data, headers):
    add_new_page(document)
    document.add_heading('Summary of Bollinger Bands, RSI, and MACD Conditions', level=1)
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False

    # Set table style
    table.style = 'Table Grid'

    # Set table header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Fill table with data
    for stock, (data, conditions, _) in stocks_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = stock
        row_cells[1].text = "Yes" if conditions[0] else "No"
        row_cells[2].text = "Yes" if conditions[1] else "No"
        row_cells[3].text = "Yes" if conditions[2] else "No"

        # Highlight cells if condition is met
        if conditions[0] and conditions[1]:
            for cell in row_cells:
                shading_element = OxmlElement('w:shd')
                shading_element.set(qn('w:fill'), 'D9EAD3')  # Light green
                cell._element.get_or_add_tcPr().append(shading_element)


# Function to create VFI and VPT table
def create_vfi_vpt_table(document, stocks_data, headers):
    add_new_page(document)
    document.add_heading('Summary of VFI and VPT Conditions', level=1)
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False

    # Set table style
    table.style = 'Table Grid'

    # Set table header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Fill table with data
    for stock, (data, conditions, _) in stocks_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = stock
        vfi_cell = row_cells[1]
        vpt_cell = row_cells[2]
        vfi_cell.text = "Yes" if conditions[3] else "No"
        vpt_cell.text = "Yes" if conditions[4] else "No"

        # Highlight cells if condition is met
        if conditions[3] and conditions[4]:
            for cell in row_cells:
                shading_element = OxmlElement('w:shd')
                shading_element.set(qn('w:fill'), 'D9EAD3')  # Light green
                cell._element.get_or_add_tcPr().append(shading_element)


# Function to create financial metrics table
def create_financial_metrics_table(document, stocks_data, headers):
    add_new_page(document)
    document.add_heading('Summary of Financial Metrics', level=1)
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False

    # Set table style
    table.style = 'Table Grid'

    # Set table header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Fill table with data
    for stock, (_, _, financial_metrics) in stocks_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = stock
        for i, header in enumerate(headers[1:], 1):
            value = financial_metrics.get(header)
            if value is None:
                cell_text = "N/A"
            elif isinstance(value, (int, float, np.floating, np.integer)):
                # Special scaling for Market Cap to keep numbers readable
                if header == "Market Cap":
                    if value >= 1e9:
                        cell_text = f"{value / 1e9:.2f}B"
                    elif value >= 1e6:
                        cell_text = f"{value / 1e6:.2f}M"
                    else:
                        cell_text = f"{value:.2f}"
                else:
                    cell_text = f"{value:.2f}"
            else:
                cell_text = str(value)
            row_cells[i].text = cell_text


# Function to fetch sector PE dynamically
def fetch_sector_pe(stock):
    ticker = yf.Ticker(stock)
    sector = ticker.info.get('sector', 'Technology')
    sector_etfs = {
        'Technology': 'XLK',
        'Healthcare': 'XLV',
        'Financial Services': 'XLF',
        'Consumer Cyclical': 'XLY',
        'Consumer Defensive': 'XLP',
        'Industrials': 'XLI',
        'Utilities': 'XLU',
        'Basic Materials': 'XLB',
        'Energy': 'XLE',
        'Real Estate': 'XLRE',
        'Communication Services': 'XLC'
    }
    sector_ticker = yf.Ticker(sector_etfs.get(sector, 'XLK'))
    sector_pe = sector_ticker.info.get('trailingPE', None)
    return sector_pe


# Function to generate and save charts
def generate_charts(stock, data):
    fig, axes = plt.subplots(5, 1, figsize=(20, 20), sharex=True)

    # Bollinger Bands
    axes[0].plot(data.index, data['Close'], label='Close Price', color='blue' )
    axes[0].plot(data.index, data['BB_High'], label='BB High', color='red', linestyle='--')
    axes[0].plot(data.index, data['BB_Low'], label='BB Low', color='green', linestyle='--')
    axes[0].set_title(f'{stock} - Close Price and Bollinger Bands')
    axes[0].legend()

    # RSI
    axes[1].plot(data.index, data['RSI'], label='RSI', color='blue')
    axes[1].axhline(30, color='green', linestyle='--')
    axes[1].axhline(70, color='red', linestyle='--')
    axes[1].set_title(f'{stock} - RSI')
    axes[1].legend()

    # MACD
    axes[2].plot(data.index, data['MACD'], label='MACD', color='blue')
    axes[2].plot(data.index, data['MACD_Signal'], label='MACD Signal', color='red')
    axes[2].set_title(f'{stock} - MACD')
    axes[2].legend()

    # VFI
    axes[3].plot(data.index, data['VFI'], label='VFI')
    axes[3].set_title(f'{stock} - VFI')
    axes[3].legend()

    # VPT
    axes[4].plot(data.index, data['VPT'], label='VPT')
    axes[4].axhline(data['VPT_POC_High'].max(), color='red', linestyle='--', label='POC High')
    axes[4].axhline(data['VPT_POC_Low'].min(), color='green', linestyle='--', label='POC Low')
    axes[4].set_title(f'{stock} - Volume Price Trend with POC')
    axes[4].legend()

    plt.tight_layout()
    plt.savefig(f'{stock}_charts.png')
    plt.close()


# Delay between requests to avoid Yahoo Finance rate limiting (seconds)
DOWNLOAD_DELAY = 2.0
RATE_LIMIT_WAIT = 60  # seconds to wait when rate limited before retry
MAX_RETRIES = 2


# Function to fetch data and calculate indicators
def check_stocks(stocks):
    stocks_data = {}

    for i, stock in enumerate(stocks):
        try:
            # Space out requests to avoid rate limiting
            if i > 0:
                time.sleep(DOWNLOAD_DELAY)

            # Fetch historical data for 1-hour interval (with retry when empty/rate limited)
            data = None
            for attempt in range(MAX_RETRIES + 1):
                try:
                    data = yf.download(stock, period="6mo", interval="1h", progress=False, auto_adjust=False)
                except Exception as dl_e:
                    err_msg = str(dl_e).lower()
                    if "rate limit" in err_msg or "too many requests" in err_msg:
                        if attempt < MAX_RETRIES:
                            print(f"[wait] Rate limited on {stock}, waiting {RATE_LIMIT_WAIT}s before retry ({attempt + 1}/{MAX_RETRIES})...")
                            time.sleep(RATE_LIMIT_WAIT)
                        else:
                            print(f"[!] Skipping {stock} - Rate limited after {MAX_RETRIES + 1} attempts.")
                            data = pd.DataFrame()
                            break
                    else:
                        raise
                # Yahoo often returns empty DataFrame on rate limit instead of raising
                if data is not None and not data.empty:
                    break
                if attempt < MAX_RETRIES:
                    print(f"[wait] No data for {stock} (possible rate limit), waiting {RATE_LIMIT_WAIT}s before retry ({attempt + 1}/{MAX_RETRIES})...")
                    time.sleep(RATE_LIMIT_WAIT)

            if data is None or data.empty:
                if data is None:
                    print(f"[!] Skipping {stock} - No data returned.")
                else:
                    print(f"[!] Skipping {stock} - No data found.")
                continue

            # Flatten MultiIndex columns if necessary
            if isinstance(data.columns, pd.MultiIndex):
                data.columns = data.columns.get_level_values(1)

            # Fix incorrect column names if all columns match the stock ticker
            expected_cols = ['Open', 'High', 'Low', 'Close', 'Volume']

            if all(col == stock for col in data.columns):  # If all columns are the stock ticker name
                if len(data.columns) == 6:  # yfinance 1.x often returns Adj Close as 6th column
                    data = data.iloc[:, :-1]  # Drop the last column (Adj Close)
                if len(data.columns) == 5:  # Now safe to rename
                    data.columns = expected_cols
                else:
                    print(f"[!] Skipping {stock} - Unexpected number of columns: {list(data.columns)}")
                    continue

            # Ensure index is datetime with timezone removed
            data.index = pd.to_datetime(data.index).tz_localize(None)

            # Resample data to 4-hour intervals
            data_4h = data.resample('4H').agg({
                'Open': 'first',
                'High': 'max',
                'Low': 'min',
                'Close': 'last',
                'Volume': 'sum'
            }).dropna()

            # Calculate Bollinger Bands with a window of 100 and 2 standard deviations
            bb_indicator = BollingerBands(close=data_4h['Close'], window=200, window_dev=2)
            data_4h['BB_High'] = bb_indicator.bollinger_hband()
            data_4h['BB_Low'] = bb_indicator.bollinger_lband()

            # Calculate RSI with a window of 10
            data_4h['RSI'] = RSIIndicator(data_4h['Close'], window=10).rsi()

            # Calculate MACD (fast 8, slow 21, signal 5)
            macd_indicator = MACD(data_4h['Close'], window_slow=21, window_fast=8, window_sign=5)
            data_4h['MACD'] = macd_indicator.macd()
            data_4h['MACD_Signal'] = macd_indicator.macd_signal()

            # Calculate VFI
            data_4h['VFI'] = calculate_vfi(data_4h['Close'], data_4h['High'], data_4h['Low'], data_4h['Volume'])

            # Calculate VPT
            data_4h['VPT'] = VolumePriceTrendIndicator(data_4h['Close'], data_4h['Volume']).volume_price_trend()
            data_4h['VPT_POC_High'] = data_4h['VPT'].rolling(window=len(data_4h)).max()
            data_4h['VPT_POC_Low'] = data_4h['VPT'].rolling(window=len(data_4h)).min()

            # Check conditions
            conditions = [
                data_4h['Close'].iloc[-1] < data_4h['BB_Low'].iloc[-1],
                data_4h['RSI'].iloc[-1] < 30,
                data_4h['MACD'].iloc[-1] > data_4h['MACD_Signal'].iloc[-1],
                data_4h['VFI'].iloc[-1] > 0,
                data_4h['VPT'].iloc[-1] > data_4h['VPT'].iloc[-2]
            ]

            # Fetch financial metrics
            ticker = yf.Ticker(stock)

            # Handle potential errors when fetching financial data
            try:
                financials = ticker.financials.T if isinstance(ticker.financials, pd.DataFrame) else pd.DataFrame()
                balance_sheet = ticker.balance_sheet.T
                cashflow = ticker.cashflow.T
            except Exception as e:
                print(f"[!] Error fetching financial data for {stock}: {e}")
                continue

            # Sort statements by period (oldest -> newest) for consistent multi-year calculations
            if not financials.empty:
                financials = financials.sort_index()
            if not balance_sheet.empty:
                balance_sheet = balance_sheet.sort_index()
            if not cashflow.empty:
                cashflow = cashflow.sort_index()

            # Load static info dict once (used for market cap, dividend rate, etc.)
            info = {}
            try:
                info = ticker.info or {}
            except Exception:
                info = {}

            # Try to get market cap (prefer fast_info, fall back to info)
            market_cap = None
            try:
                fast_info = getattr(ticker, "fast_info", None)
                market_cap = getattr(fast_info, "market_cap", None) if fast_info is not None else None
            except Exception:
                market_cap = None
            if market_cap is None:
                market_cap = info.get('marketCap', None)

            # Initialize metrics with the new intent
            financial_metrics = {
                'P/E (3Y Avg)': None,
                'P/E (1Y)': None,
                'Revenue Growth (2Y %)': None,
                'Profit Growth (2Y vs Today)': None,
                'Shares Outstanding (2Y)': None,
                'Div Rate (FWD)': None,
                'FCF Multiple (MC/FCF)': None,
                'Market Cap': market_cap,
                'Sector PE': fetch_sector_pe(stock),
            }

            # Net income based metrics (P/E and profit growth)
            if not financials.empty and 'Net Income' in financials.columns:
                ni = financials['Net Income'].dropna()

                # 1-year P/E = Market Cap / last year's profit
                if len(ni) >= 1 and market_cap is not None:
                    last_profit = ni.iloc[-1]
                    if pd.notnull(last_profit) and last_profit != 0:
                        financial_metrics['P/E (1Y)'] = market_cap / last_profit

                # 3-year average P/E and profit growth vs 2 years ago
                if len(ni) >= 3:
                    if market_cap is not None:
                        avg_profit = ni.iloc[-3:].mean()
                        if pd.notnull(avg_profit) and avg_profit != 0:
                            financial_metrics['P/E (3Y Avg)'] = market_cap / avg_profit

                    latest_profit = ni.iloc[-1]
                    profit_two_years_ago = ni.iloc[-3]
                    if pd.notnull(latest_profit) and pd.notnull(profit_two_years_ago):
                        if latest_profit > profit_two_years_ago:
                            financial_metrics['Profit Growth (2Y vs Today)'] = "Up"
                        elif latest_profit < profit_two_years_ago:
                            financial_metrics['Profit Growth (2Y vs Today)'] = "Down"
                        else:
                            financial_metrics['Profit Growth (2Y vs Today)'] = "Flat"

            # Revenue growth: today vs 2 years ago
            if not financials.empty and 'Total Revenue' in financials.columns:
                rev = financials['Total Revenue'].dropna()
                if len(rev) >= 3:
                    latest_rev = rev.iloc[-1]
                    rev_two_years_ago = rev.iloc[-3]
                    if pd.notnull(latest_rev) and pd.notnull(rev_two_years_ago) and rev_two_years_ago != 0:
                        growth = (latest_rev - rev_two_years_ago) / abs(rev_two_years_ago) * 100
                        financial_metrics['Revenue Growth (2Y %)'] = growth

            # Shares outstanding: buyback / issuing / neutral over last 2 years
            if not balance_sheet.empty and 'Share Issued' in balance_sheet.columns:
                shares = balance_sheet['Share Issued'].dropna()
                if len(shares) >= 3:
                    latest_shares = shares.iloc[-1]
                    shares_two_years_ago = shares.iloc[-3]
                    if pd.notnull(latest_shares) and pd.notnull(shares_two_years_ago) and shares_two_years_ago != 0:
                        change_pct = (latest_shares - shares_two_years_ago) / shares_two_years_ago * 100
                        flag = "Neutral"
                        if change_pct > 5:
                            flag = "Issuing Shares"
                        elif change_pct < -5:
                            flag = "Buying Back"
                        financial_metrics['Shares Outstanding (2Y)'] = f"{flag} ({change_pct:.1f}%)"

            # Free cash flow multiple: Market Cap / (Cash from operations - capital expenditures) for last year
            if not cashflow.empty and market_cap is not None:
                cf_row = cashflow.iloc[-1]

                # Cash from operations – try several common Yahoo labels
                cfo = cf_row.get('Total Cash From Operating Activities')
                if pd.isnull(cfo):
                    cfo = cf_row.get('Net Cash Provided by Operating Activities')
                if pd.isnull(cfo):
                    cfo = cf_row.get('Net Cash Provided by (Used in) Operating Activities')

                # Capital expenditures – try common variants
                capex = cf_row.get('Capital Expenditures')
                if pd.isnull(capex):
                    capex = cf_row.get('Purchase Of Property Plant And Equipment')

                if pd.notnull(cfo) and pd.notnull(capex):
                    fcf = cfo - capex
                    if fcf not in (0, None):
                        financial_metrics['FCF Multiple (MC/FCF)'] = market_cap / fcf

            # Forward dividend rate (per share, annualised, when available)
            div_rate_fwd = info.get('dividendRate')
            if pd.notnull(div_rate_fwd):
                financial_metrics['Div Rate (FWD)'] = div_rate_fwd

            stocks_data[stock] = (data_4h, conditions, financial_metrics)

        except Exception as e:
            print(f"[!] Skipping {stock} due to error: {e}")
            continue  # Move to the next stock if any issue occurs

    return stocks_data


# Updated main function for creating the report
def run_check_once(stocks):
    document = Document()
    document.add_heading('Stock Analysis Report', 0)
    document.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    stocks_data = check_stocks(stocks)

    if not stocks_data:
        print("No stock data was retrieved (all downloads failed or were skipped, e.g. due to rate limiting).")
        print("Try again later or use a shorter list of symbols. Report not generated.")
        return

    create_summary_table(document, stocks_data, ["Stock Symbol", "Below BB Low", "RSI Extreme",
                                                 "MACD Crossover"])
    create_vfi_vpt_table(document, stocks_data, ["Stock Symbol", "VFI Positive", "VPT Increasing"])
    create_financial_metrics_table(
        document,
        stocks_data,
        [
            "Stock Symbol",
            "P/E (3Y Avg)",
            "P/E (1Y)",
            "Revenue Growth (2Y %)",
            "Profit Growth (2Y vs Today)",
            "Shares Outstanding (2Y)",
            "Div Rate (FWD)",
            "FCF Multiple (MC/FCF)",
            "Market Cap",
            "Sector PE",
        ],
    )

    document.add_heading('Stock Charts', level=1)
    (document.add_paragraph
     ('The following charts provide a visual representation of the stock\'s performance, '
      'Bollinger Bands, RSI, and MACD indicators.'))

    for stock, (data, conditions, _) in stocks_data.items():
        if sum(conditions[:3])>=2:
            add_new_page(document)
            document.add_heading(f'{stock} Charts', level=2)
            generate_charts(stock, data)
            document.add_picture(f'{stock}_charts.png', width=Inches(6))
            os.remove(f'{stock}_charts.png')

    document.add_heading('Explanation', level=1)
    document.add_paragraph(
        'This report provides a summary of several technical and financial metrics for the specified stocks.\n\n'
        'Tables Summary:\n'
        '1. **Summary of Bollinger Bands, RSI, and MACD Conditions**: This table shows whether the stock\'s price is below the lower Bollinger Band, '
        'whether the RSI indicates an extreme condition (overbought/oversold), and whether there is a bullish MACD crossover.\n'
        '2. **Summary of VFI and VPT Conditions**: This table indicates if the Volume Flow Indicator (VFI) is positive and if the Volume Price Trend (VPT) '
        'is increasing, suggesting more buying activity.\n'
        '3. **Summary of Financial Metrics**:\n'
        '   - **P/E (3Y Avg)**: Market capitalization divided by the average Net Income over the last three reported years.\n'
        '   - **P/E (1Y)**: Market capitalization divided by the most recent year\'s Net Income.\n'
        '   - **Revenue Growth (2Y %)**: Percentage change in Total Revenue from two years ago to the most recent year.\n'
        '   - **Profit Growth (2Y vs Today)**: Indicates whether Net Income today is Up, Down, or Flat compared to two years ago.\n'
        '   - **Shares Outstanding (2Y)**: Flags whether the company has been Issuing Shares, Buying Back shares, or is Neutral based on the percentage change in '
        'shares issued over the last two years (threshold ±5%).\n'
        '   - **Div Rate (FWD)**: Forward annual dividend rate per share, based on the company\'s expected dividend payments.\n'
        '   - **FCF Multiple (MC/FCF)**: Market capitalization divided by Free Cash Flow, where Free Cash Flow is calculated as Cash From Operations minus Capital Expenditures '
        'in the most recent year.\n'
        '   - **Market Cap**: Current equity market capitalization of the company (share price multiplied by current shares outstanding), shown in billions (B) or millions (M) '
        'with two decimal places.\n'
        '   - **Sector PE**: Trailing P/E ratio for a representative sector ETF corresponding to the company\'s primary sector (e.g., XLV for Healthcare, XLK for Technology).\n\n'
        'Stock Charts Explanation:\n'
        'The charts for each stock include three sections:\n'
        '1. **Close Price and Bollinger Bands**: Shows the stock\'s close price along with the Bollinger Bands to indicate potential overbought or oversold conditions.\n'
        '2. **RSI**: Displays the Relative Strength Index to highlight potential overbought or oversold levels (above 70 or below 30, respectively).\n'
        '3. **MACD**: Shows the MACD line and signal line to identify potential bullish or bearish crossovers.\n\n'
        'Reading the Charts:\n'
        'For each stock, observe the close price in relation to the Bollinger Bands. A price below the lower band could indicate a buying opportunity. '
        'Check the RSI values to see if the stock is overbought or oversold. Finally, look at the MACD for potential bullish crossovers where the MACD line crosses '
        'above the signal line, indicating potential upward momentum.'
    )

    # Save the document
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    document.save(f'Stock_Analysis_Report_{now}.docx')
    print(f"Stock analysis report saved to 'Stock_Analysis_Report_{now}.docx'.")


if __name__ == "__main__":
    custom_stocks_input = input("Enter a comma-separated list of stock symbols (or press Enter to use NASDAQ-100): ")
    if custom_stocks_input:
        custom_stocks = [s.strip() for s in custom_stocks_input.split(',')]
        run_check_once(custom_stocks)
    else:
        nasdaq100_df = pd.read_csv('stockList.csv')
        nasdaq100_stocks = nasdaq100_df['Symbol'].tolist()
        run_check_once(nasdaq100_stocks)
